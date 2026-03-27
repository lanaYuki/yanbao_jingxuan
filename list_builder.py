"""
list_builder.py - 生成研报 list docx（研报清单）
格式参考：研报模版.docx 最后几页
- 无边框两列表格：日期列（约1501 dxa）+ 标题列（约7911 dxa）
- 字体：仿宋_GB2312 13pt；英文/数字 Times New Roman 13pt（混排）
- 分类标题：黑体，加粗，下划线，段前段后间距
- 行高：atLeast 340 twip，cantSplit，不分页
- 日期格式：YYYY.M.D（无补零）
"""

import re
import os
import zipfile
from io import BytesIO
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

FONT_FANG_SONG = '仿宋_GB2312'
FONT_TIMES = 'Times New Roman'
FONT_HEITI = '黑体'
SIZE_BODY = Pt(13)         # sz=26 half-points
SIZE_HEADING = Pt(16)      # 分类标题

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'CGI每周研报精选（第259期）-摘要.docx')

# 4个固定分类（顺序固定）
CATEGORIES = ['宏观', '策略及大宗商品', '固定收益', '行业']

# 表格列宽（twips/dxa）
COL_DATE_W  = 1501
COL_TITLE_W = 7911


# ── 解析工具 ──────────────────────────────────────────────────


def parse_pasted_text(text: str) -> list:
    """
    从中金点睛粘贴文本中提取研报列表。

    中金点睛的研报条目结构（每条之间紧连）：
      收藏  [可选标签1] [可选标签2]
      研报标题（第一个长句，通常 > 8 个字）
      [可选：副标题/摘要行]
      元信息行：作者 | N页 | X小时前/X天前 | 分类

    策略：
    - 先按"收藏"拆分各条目（每条以"收藏"开头，或第一条在"收藏"之前）
    - 每条内，跳过"收藏"行本身及其后的短标签行（≤8字且无中文停顿标点）
    - 取第一个长行（> 8 字）作为标题
    - 从最后一行（元信息行，含"|"）提取日期

    返回 [{'title': str, 'date': str}, ...]
    """
    results = []
    lines = [l.strip() for l in text.strip().splitlines()]
    lines = [l for l in lines if l]

    # 按"收藏"分割条目
    blocks = []
    current = []
    for line in lines:
        if line == '收藏' and current:
            blocks.append(current)
            current = ['收藏']
        else:
            current.append(line)
    if current:
        blocks.append(current)

    # 如果没有"收藏"关键词，退回整体当一块处理
    if not blocks or (len(blocks) == 1 and blocks[0][0] != '收藏'):
        blocks = [lines]

    for block in blocks:
        title, date = _extract_title_date(block)
        if title:
            results.append({'title': title, 'date': date})

    return results


def _is_tag_line(line: str) -> bool:
    """
    判断是否为标签行（短词，如"固收+"、"宏观"、"策略"、"信用"、"收藏"等）。
    标签行特征：长度 ≤ 8 字，且不含任何中文标点或 | 符号，不含数字+前等时间词。
    """
    if len(line) > 10:
        return False
    # 含 | 的是元信息行
    if '|' in line:
        return False
    # 含时间词的是元信息行
    if re.search(r'\d+\s*(?:分钟前|小时前|天前)', line):
        return False
    # 含句号、逗号、冒号等中文停顿标点，或英文标点，通常是正文行
    if re.search(r'[，。：；「」（）【】、]', line):
        return False
    return True


def _extract_title_date(block: list) -> tuple:
    """
    从一个条目的行列表中提取 (title, date)。
    """
    meta_line = ''
    title = ''

    # 找最后一个含"|"的行作为元信息行
    for line in reversed(block):
        if '|' in line and re.search(r'(?:页|分钟前|小时前|天前|\d{4})', line):
            meta_line = line
            break

    # 从 block 中找标题：跳过"收藏"行和标签行，取第一个长行
    for line in block:
        if line == '收藏':
            continue
        if _is_tag_line(line):
            continue
        if line == meta_line:
            continue
        # 第一个非标签、非元信息的行就是标题
        title = line
        break

    date = _parse_date_from_meta(meta_line) if meta_line else _today_date()
    return title, date


def _today_date() -> str:
    today = datetime.now()
    return f'{today.year}.{today.month}.{today.day}'


def _parse_date_from_meta(meta_line: str) -> str:
    """
    从元信息行提取日期，格式返回 YYYY.M.D（无补零）。
    优先匹配绝对日期，否则将相对时间转换为今天。
    """
    today = datetime.now()

    # 绝对日期（yyyy.mm.dd / yyyy-mm-dd / yyyy年mm月dd日）
    m = re.search(r'(\d{4})[.\-年](\d{1,2})[.\-月](\d{1,2})', meta_line)
    if m:
        return f'{m.group(1)}.{int(m.group(2))}.{int(m.group(3))}'

    # 相对时间
    m = re.search(r'(\d+)\s*天前', meta_line)
    if m:
        d = today - timedelta(days=int(m.group(1)))
        return f'{d.year}.{d.month}.{d.day}'

    m = re.search(r'(\d+)\s*小时前', meta_line)
    if m:
        d = today - timedelta(hours=int(m.group(1)))
        return f'{d.year}.{d.month}.{d.day}'

    m = re.search(r'(\d+)\s*分钟前', meta_line)
    if m:
        return f'{today.year}.{today.month}.{today.day}'

    return f'{today.year}.{today.month}.{today.day}'


# ── 文档构建工具 ─────────────────────────────────────────────


def _is_ascii_char(ch: str) -> bool:
    return ch.isascii() and (ch.isalpha() or ch.isdigit())


def _set_run_fonts_size(rPr, cn_font: str, en_font: str, size_pt: float):
    """在已有 rPr 上设置字体和字号。"""
    sz_val = str(int(size_pt * 2))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)
    for tag in ('w:sz', 'w:szCs'):
        el = rPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rPr.append(el)
        el.set(qn('w:val'), sz_val)


def _add_mixed_runs_to_para(para, text: str, size_pt: float = 13,
                             bold: bool = False, color: str = None):
    """
    向段落中写入中英混排文本：中文用仿宋_GB2312，英文/数字用 Times New Roman。
    """
    if not text:
        return

    # 按中/英文分段
    segments = []
    current = text[0]
    current_is_en = _is_ascii_char(text[0])
    for ch in text[1:]:
        ch_is_en = _is_ascii_char(ch)
        if ch_is_en == current_is_en:
            current += ch
        else:
            segments.append((current, current_is_en))
            current = ch
            current_is_en = ch_is_en
    segments.append((current, current_is_en))

    for seg_text, is_en in segments:
        run = para.add_run(seg_text)
        run.bold = bold
        rpr = run._r.get_or_add_rPr()
        en_font = FONT_TIMES if is_en else FONT_FANG_SONG
        cn_font = FONT_FANG_SONG
        _set_run_fonts_size(rpr, cn_font, en_font, size_pt)
        if color:
            col_el = rpr.find(qn('w:color'))
            if col_el is None:
                col_el = OxmlElement('w:color')
                rpr.append(col_el)
            col_el.set(qn('w:val'), color)


def _set_cell_para_format(para):
    """表格单元格段落格式：行距auto单倍，无段前段后，无首行缩进。"""
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # 消除首行缩进
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')
    # snapToGrid=0, adjustRightInd=0
    for tag, val in [('w:snapToGrid', '0'), ('w:adjustRightInd', '0')]:
        el = pPr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            pPr.append(el)
        el.set(qn('w:val'), val)


def _set_row_props(row):
    """行属性：cantSplit，最小行高340 twip。"""
    trPr = row._tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        row._tr.insert(0, trPr)
    # cantSplit
    cs = trPr.find(qn('w:cantSplit'))
    if cs is None:
        cs = OxmlElement('w:cantSplit')
        trPr.append(cs)
    # trHeight atLeast 340
    trH = trPr.find(qn('w:trHeight'))
    if trH is None:
        trH = OxmlElement('w:trHeight')
        trPr.append(trH)
    trH.set(qn('w:val'), '340')
    trH.set(qn('w:hRule'), 'atLeast')


def _set_cell_width(cell, width_dxa: int):
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._tc.insert(0, tcPr)
    tcW = tcPr.find(qn('w:tcW'))
    if tcW is None:
        tcW = OxmlElement('w:tcW')
        tcPr.append(tcW)
    tcW.set(qn('w:w'), str(width_dxa))
    tcW.set(qn('w:type'), 'dxa')


def _set_cell_vAlign_center(cell):
    tcPr = cell._tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        cell._tc.insert(0, tcPr)
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = OxmlElement('w:vAlign')
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), 'center')


def _set_table_no_borders(table):
    """将表格所有边框设为 none，并去掉单元格间距。"""
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)

    # 移除已有 tblBorders（如果有）
    old_bdr = tblPr.find(qn('w:tblBorders'))
    if old_bdr is not None:
        tblPr.remove(old_bdr)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

    # 表格宽度 autofit
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')

    # 单元格 margin
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is None:
        tblCellMar = OxmlElement('w:tblCellMar')
        tblPr.append(tblCellMar)
    for side, val in [('top', '0'), ('left', '108'), ('bottom', '0'), ('right', '108')]:
        el = tblCellMar.find(qn(f'w:{side}'))
        if el is None:
            el = OxmlElement(f'w:{side}')
            tblCellMar.append(el)
        el.set(qn('w:w'), val)
        el.set(qn('w:type'), 'dxa')


def _add_section_heading_para(doc: Document, text: str):
    """
    分类标题段落：黑体 16pt，加粗，下划线，段前段后各 156 twip（约8pt）。
    """
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_before = Twips(156)
    pf.space_after = Twips(156)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.0
    # 无首行缩进
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')

    run = para.add_run(text)
    run.bold = True
    run.underline = True
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_HEITI)
    rFonts.set(qn('w:hAnsi'), FONT_HEITI)
    rFonts.set(qn('w:eastAsia'), FONT_HEITI)
    sz_val = str(int(SIZE_HEADING.pt * 2))
    for tag in ('w:sz', 'w:szCs'):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn('w:val'), sz_val)


def _add_page_break(doc: Document):
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)


def _add_report_table(doc: Document, items: list):
    """
    添加一个研报清单表格：2列（日期 | 标题），无边框。
    items: [{'title': str, 'date': str}, ...]
    """
    table = doc.add_table(rows=len(items), cols=2)
    _set_table_no_borders(table)

    for row_idx, item in enumerate(items):
        row = table.rows[row_idx]
        _set_row_props(row)

        # ── 日期列 ──
        date_cell = row.cells[0]
        _set_cell_width(date_cell, COL_DATE_W)
        # 清除默认段落
        for p in date_cell.paragraphs:
            for r in p.runs:
                r.text = ''
        date_para = date_cell.paragraphs[0]
        _set_cell_para_format(date_para)
        _add_mixed_runs_to_para(date_para, item.get('date', ''), size_pt=13)

        # ── 标题列 ──
        title_cell = row.cells[1]
        _set_cell_width(title_cell, COL_TITLE_W)
        _set_cell_vAlign_center(title_cell)
        for p in title_cell.paragraphs:
            for r in p.runs:
                r.text = ''
        title_para = title_cell.paragraphs[0]
        _set_cell_para_format(title_para)
        _add_mixed_runs_to_para(title_para, item.get('title', ''), size_pt=13)


# ── 主构建函数 ────────────────────────────────────────────────


def _sort_items_by_date(items: list) -> list:
    """按日期升序排序（日期早的在先）。日期格式 YYYY.M.D，无法解析的排最后。"""
    def _to_date(item):
        try:
            parts = item.get('date', '').split('.')
            return datetime(int(parts[0]), int(parts[1]), int(parts[2]))
        except Exception:
            return datetime.max
    return sorted(items, key=_to_date)


def append_list_to_doc(doc: Document, category_data: dict):
    """
    将研报清单内容追加到已有 doc 对象末尾（供 build_jingxuan 合并调用）。
    在清单首页前插入分页符，然后写入"附：一周主要研报回顾"标题和各分类表格。
    """
    _add_page_break(doc)

    # 总标题
    title_para = doc.add_paragraph()
    pf = title_para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Twips(156)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.0
    pPr = title_para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')
    run = title_para.add_run('附：一周主要研报回顾')
    run.bold = True
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_HEITI)
    rFonts.set(qn('w:hAnsi'), FONT_HEITI)
    rFonts.set(qn('w:eastAsia'), FONT_HEITI)
    sz_val = str(int(Pt(16).pt * 2))
    for tag in ('w:sz', 'w:szCs'):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn('w:val'), sz_val)

    active_cats = [cat for cat in CATEGORIES if category_data.get(cat)]
    last_cat = active_cats[-1] if active_cats else None

    for cat_i, cat in enumerate(CATEGORIES):
        items = category_data.get(cat, [])
        if not items:
            continue
        if cat == '行业' and cat == last_cat and cat_i > 0:
            _add_page_break(doc)
        _add_section_heading_para(doc, cat)
        _add_report_table(doc, _sort_items_by_date(items))


def build_list_docx(category_data: dict, output_path: str, issue: str = ''):
    """独立构建研报清单 docx（保留备用）。"""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin = Cm(3.2)
        sec.bottom_margin = Cm(2.2)
        sec.left_margin = Cm(2.2)
        sec.right_margin = Cm(2.2)
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)
    append_list_to_doc(doc, category_data)
    doc.save(output_path)

    # 添加页眉
    _add_header_to_list_docx(output_path)


def _add_header_to_list_docx(output_path: str):
    """从模板复制 header4（default header）到 list docx。"""
    if not os.path.exists(TEMPLATE_PATH):
        return

    REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
    W_NS_LOCAL = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R_NS_LOCAL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
    A_NS_H = 'http://schemas.openxmlformats.org/drawingml/2006/main'
    WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    MOVE_DOWN_EMU = 120000

    try:
        with zipfile.ZipFile(TEMPLATE_PATH, 'r') as tmpl:
            header4_xml = tmpl.read('word/header4.xml')
            header4_rels = tmpl.read('word/_rels/header4.xml.rels')
            image2_data = tmpl.read('word/media/image2.png')
    except Exception:
        return

    h4_tree = etree.fromstring(header4_xml)
    for wsp in h4_tree.iter(f'{{{WPS_NS}}}wsp'):
        for xfrm in wsp.iter(f'{{{A_NS_H}}}xfrm'):
            off = xfrm.find(f'{{{A_NS_H}}}off')
            ext = xfrm.find(f'{{{A_NS_H}}}ext')
            if off is not None and ext is not None:
                ext.set('cx', '6726754')
                off.set('x', '0')
                off.set('y', str(2050000 + MOVE_DOWN_EMU))
            break

    for xfrm in h4_tree.xpath("//a:xfrm", namespaces={'a': A_NS_H}):
        if f'{{{WPS_NS}}}wsp' not in [anc.tag for anc in xfrm.iterancestors()]:
            off = xfrm.find(f'{{{A_NS_H}}}off')
            if off is not None:
                current_y = int(off.get('y', '0'))
                off.set('y', str(current_y + MOVE_DOWN_EMU))

    for txbx_para in h4_tree.findall(f'.//{{{W_NS_LOCAL}}}txbxContent/{{{W_NS_LOCAL}}}p'):
        pPr = txbx_para.find(f'{{{W_NS_LOCAL}}}pPr')
        if pPr is None:
            pPr = etree.SubElement(txbx_para, f'{{{W_NS_LOCAL}}}pPr')
            txbx_para.insert(0, pPr)
        jc = pPr.find(f'{{{W_NS_LOCAL}}}jc')
        if jc is None:
            jc = etree.SubElement(pPr, f'{{{W_NS_LOCAL}}}jc')
        jc.set(f'{{{W_NS_LOCAL}}}val', 'right')
        ind = pPr.find(f'{{{W_NS_LOCAL}}}ind')
        if ind is not None:
            pPr.remove(ind)
        new_ind = etree.SubElement(pPr, f'{{{W_NS_LOCAL}}}ind')
        new_ind.set(f'{{{W_NS_LOCAL}}}left', '0')
        new_ind.set(f'{{{W_NS_LOCAL}}}firstLine', '0')
        spacing = pPr.find(f'{{{W_NS_LOCAL}}}spacing')
        if spacing is None:
            spacing = etree.SubElement(pPr, f'{{{W_NS_LOCAL}}}spacing')
        spacing.set(f'{{{W_NS_LOCAL}}}before', '0')
        spacing.set(f'{{{W_NS_LOCAL}}}after', '0')
        for run in txbx_para.findall(f'{{{W_NS_LOCAL}}}r'):
            rPr = run.find(f'{{{W_NS_LOCAL}}}rPr')
            if rPr is None:
                rPr = etree.SubElement(run, f'{{{W_NS_LOCAL}}}rPr')
            sz = rPr.find(f'{{{W_NS_LOCAL}}}sz')
            if sz is None:
                sz = etree.SubElement(rPr, f'{{{W_NS_LOCAL}}}sz')
            sz.set(f'{{{W_NS_LOCAL}}}val', '28')

    header4_xml_out = etree.tostring(h4_tree, xml_declaration=True, encoding='UTF-8', standalone=True)

    buf = BytesIO()
    with zipfile.ZipFile(output_path, 'r') as zin:
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == 'word/_rels/document.xml.rels':
                    tree = etree.fromstring(data)
                    rel = etree.SubElement(tree, f'{{{REL_NS}}}Relationship')
                    rel.set('Id', 'rId101')
                    rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header')
                    rel.set('Target', 'header4.xml')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                elif item.filename == 'word/document.xml':
                    tree = etree.fromstring(data)
                    sectPr = tree.find(f'.//{{{W_NS_LOCAL}}}sectPr')
                    if sectPr is not None:
                        for old in sectPr.findall(f'{{{W_NS_LOCAL}}}headerReference'):
                            sectPr.remove(old)
                        h4 = etree.Element(f'{{{W_NS_LOCAL}}}headerReference')
                        h4.set(f'{{{W_NS_LOCAL}}}type', 'default')
                        h4.set(f'{{{R_NS_LOCAL}}}id', 'rId101')
                        sectPr.insert(0, h4)
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                elif item.filename == '[Content_Types].xml':
                    tree = etree.fromstring(data)
                    ov = etree.SubElement(tree, f'{{{CT_NS}}}Override')
                    ov.set('PartName', '/word/header4.xml')
                    ov.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                zout.writestr(item, data)

            zout.writestr('word/header4.xml', header4_xml_out)
            zout.writestr('word/_rels/header4.xml.rels', header4_rels)
            zout.writestr('word/media/image2.png', image2_data)

    with open(output_path, 'wb') as f:
        f.write(buf.getvalue())
