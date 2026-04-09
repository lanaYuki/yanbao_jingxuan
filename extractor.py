"""
extractor.py - 从研报源 docx 中提取所需内容
使用 lxml 直接解析 XML 以正确处理合并单元格和高亮。
"""

import re
import zipfile

_FIGURE_RE = re.compile(r'[（(]图表\s*\d+[）)]')
from docx import Document
from lxml import etree

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def w(tag: str) -> str:
    return f'{{{W_NS}}}{tag}'


def _get_xml_tree(docx_path: str):
    with zipfile.ZipFile(docx_path) as z:
        content = z.read('word/document.xml')
    return etree.fromstring(content)


def extract_date(docx_path: str) -> str:
    """
    从研报源第0个表格的 DATE 域缓存值提取日期。
    返回格式如：2026年3月7日
    """
    tree = _get_xml_tree(docx_path)
    body = tree.find(w('body'))
    tbls = body.findall(w('tbl'))
    if not tbls:
        return ''

    tbl0 = tbls[0]
    # 找所有文本节点，日期格式为 yyyy.MM.dd
    texts = [t.text for t in tbl0.findall('.//' + w('t')) if t.text]
    for text in texts:
        m = re.search(r'\b(\d{4})\.(\d{2})\.(\d{2})\b', text)
        if m:
            year, month, day = m.group(1), str(int(m.group(2))), str(int(m.group(3)))
            return f'{year}年{month}月{day}日'
    return ''


def extract_institution(docx_path: str) -> str:
    """
    从研报源第0个表格识别发布机构。
    返回："中金公司研究部" 或 "中金研究院"
    """
    tree = _get_xml_tree(docx_path)
    body = tree.find(w('body'))
    tbls = body.findall(w('tbl'))
    if not tbls:
        return '中金公司研究部'

    tbl_text = ''.join(t.text for t in tbls[0].findall('.//' + w('t')) if t.text)
    if '中金研究院' in tbl_text:
        return '中金研究院'
    return '中金公司研究部'


def extract_title(docx_path: str) -> str:
    """
    从研报源 sdt 中提取报告标题（tag=BDCONTENTCONTROL_AUTO_SAVE_ZJTitle）。
    """
    tree = _get_xml_tree(docx_path)

    for sdt in tree.findall('.//' + w('sdt')):
        sdt_pr = sdt.find(w('sdtPr'))
        if sdt_pr is None:
            continue
        tag_el = sdt_pr.find(w('tag'))
        if tag_el is None:
            continue
        tag_val = tag_el.get(w('val'), '')
        if 'ZJTitle' in tag_val or 'Title' in tag_val:
            sdt_content = sdt.find(w('sdtContent'))
            if sdt_content is not None:
                texts = [t.text for t in sdt_content.findall('.//' + w('t')) if t.text]
                title = ''.join(texts).strip()
                if title:
                    return title
    return ''


def extract_authors(docx_path: str) -> list:
    """
    从研报源第1个表格提取作者信息。
    返回列表：[{"name": str, "sac": str, "sfc": str}, ...]
    删除邮箱，将"执证编号"统一为"执业证书编号"。
    """
    doc = Document(docx_path)
    if len(doc.tables) < 2:
        return []

    author_table = doc.tables[1]

    # 收集每列的段落文本（按列分组）
    col_data = {}
    for row in author_table.rows:
        for c_i, cell in enumerate(row.cells):
            if c_i not in col_data:
                col_data[c_i] = []
            for para in cell.paragraphs:
                t = para.text.strip()
                if t:
                    col_data[c_i].append(t)

    authors = []
    for c_i, lines in col_data.items():
        name = ''
        sac = ''
        sfc = ''
        for line in lines:
            if '分析员' in line or '联系人' in line:
                name = line.replace('分析员', '').replace('联系人', '').strip()
            elif 'SAC' in line:
                m = re.search(r'[SL]\d{13,}', line)
                if m:
                    sac = m.group(0)
            elif 'SFC' in line:
                m = re.search(r'[A-Z]{3}\d{3}', line)
                if m:
                    sfc = m.group(0)
            # 跳过邮箱
        if name and sac:
            authors.append({'name': name, 'sac': sac, 'sfc': sfc})

    return authors


def _load_style_bold_map(docx_path: str) -> dict:
    """
    解析 word/styles.xml，返回 {styleId: bool} 表示该样式是否为粗体。
    沿 basedOn 链向上继承；字体设置为黑体也视为粗体。
    """
    try:
        with zipfile.ZipFile(docx_path) as z:
            if 'word/styles.xml' not in z.namelist():
                return {}
            styles_xml = z.read('word/styles.xml')
    except Exception:
        return {}

    styles_tree = etree.fromstring(styles_xml)
    raw = {}
    for style in styles_tree.findall('.//' + w('style')):
        sid = style.get(w('styleId'))
        if not sid:
            continue
        based = style.find(w('basedOn'))
        based_id = based.get(w('val')) if based is not None else None
        rPr = style.find(w('rPr'))
        bold_val = None
        if rPr is not None:
            b_el = rPr.find(w('b'))
            if b_el is not None:
                val = b_el.get(w('val'), '1')
                bold_val = val not in ('0', 'false')
            if bold_val is None:
                rf = rPr.find(w('rFonts'))
                if rf is not None and '黑体' in rf.get(w('eastAsia'), ''):
                    bold_val = True
        raw[sid] = {'bold': bold_val, 'basedOn': based_id}

    resolved = {}

    def resolve(sid):
        if sid in resolved:
            return resolved[sid]
        if sid not in raw:
            resolved[sid] = False
            return False
        info = raw[sid]
        if info['bold'] is not None:
            resolved[sid] = info['bold']
            return info['bold']
        if info['basedOn']:
            result = resolve(info['basedOn'])
            resolved[sid] = result
            return result
        resolved[sid] = False
        return False

    for sid in raw:
        resolve(sid)
    return resolved


def extract_highlighted_paragraphs(docx_path: str) -> list:
    """
    从研报源第2个表格（正文区）提取所有黄色高亮段落。
    使用 lxml 直接操作 XML，正确处理合并单元格。
    返回列表：[{
        "runs": [{"text": str, "bold": bool, "footnote_ref": int|None}],
        "para_text": str,
        "has_triangle": bool,  # 段落是否为项目符号段落（numPr），输出时统一用 ►
    }, ...]
    footnote_ref: 该 run 后紧跟的脚注编号（源文件中），None 表示无脚注引用。
    """
    tree = _get_xml_tree(docx_path)
    body = tree.find(w('body'))
    tbls = body.findall(w('tbl'))

    if len(tbls) < 3:
        return []

    style_bold_map = _load_style_bold_map(docx_path)
    content_tbl = tbls[2]
    result = []
    seen_texts = set()

    for para in content_tbl.findall('.//' + w('p')):
        # 检查是否有 yellow 高亮
        hls = para.findall('.//' + w('highlight'))
        has_yellow = any(hl.get(w('val')) == 'yellow' for hl in hls)
        if not has_yellow:
            continue

        # 检测是否为项目符号段落（numPr），不论原符号是什么，输出时统一用 ►
        pPr = para.find(w('pPr'))
        has_triangle = (pPr is not None and pPr.find(w('numPr')) is not None)

        # 段落样式决定的默认 bold（run 无显式 w:b 时继承）
        para_style_bold = False
        if pPr is not None:
            pStyle = pPr.find(w('pStyle'))
            if pStyle is not None:
                para_style_bold = style_bold_map.get(pStyle.get(w('val'), ''), False)

        # 提取文本和run级别bold，同时捕捉脚注引用
        runs_data = []
        for child in para:
            tag = child.tag
            if tag == w('r'):
                fn_ref = child.find(w('footnoteReference'))
                if fn_ref is not None:
                    fn_id = fn_ref.get(w('id'))
                    if runs_data and fn_id is not None:
                        runs_data[-1]['footnote_ref'] = int(fn_id)
                    continue
                run_texts = [t.text for t in child.findall(w('t')) if t.text]
                if not run_texts:
                    run_texts = [t.text for t in child.findall('.//' + w('t')) if t.text]
                if not run_texts:
                    continue
                text = ''.join(run_texts)
                bold = _is_run_element_bold(child, para_style_bold)
                runs_data.append({'text': text, 'bold': bold, 'footnote_ref': None})

            elif tag == w('hyperlink'):
                for run in child.findall(w('r')):
                    fn_ref = run.find(w('footnoteReference'))
                    if fn_ref is not None:
                        fn_id = fn_ref.get(w('id'))
                        if runs_data and fn_id is not None:
                            runs_data[-1]['footnote_ref'] = int(fn_id)
                        continue
                    run_texts = [t.text for t in run.findall(w('t')) if t.text]
                    if not run_texts:
                        run_texts = [t.text for t in run.findall('.//' + w('t')) if t.text]
                    if not run_texts:
                        continue
                    text = ''.join(run_texts)
                    bold = _is_run_element_bold(run, para_style_bold)
                    runs_data.append({'text': text, 'bold': bold, 'footnote_ref': None})

        if not runs_data:
            continue

        para_text = ''.join(r['text'] for r in runs_data)
        if not para_text.strip():
            continue

        # 删除图表引用文字，如"(图表1)""（图表3）"等
        # 先对每个 run 单独替换（处理同一 run 内的情况）
        runs_data = [{**r, 'text': _FIGURE_RE.sub('', r['text'])} for r in runs_data]
        runs_data = [r for r in runs_data if r['text']]
        # 再对整段兜底（处理跨 run 的情况）：
        # 将所有字符附上 run 属性，对合并文本做替换，再按字符 run 属性重组
        joined = ''.join(r['text'] for r in runs_data)
        cleaned = _FIGURE_RE.sub('', joined)
        if cleaned != joined:
            # 构建字符→run属性映射
            char_attrs = []  # [(bold, footnote_ref), ...]
            for r in runs_data:
                for _ in r['text']:
                    char_attrs.append((r['bold'], r.get('footnote_ref')))
            # 找出被删除的字符位置（对应 joined 中哪些位置被删掉）
            # 用替换后位置映射重建 runs
            kept_indices = []
            ci = 0
            for m in _FIGURE_RE.finditer(joined):
                kept_indices.extend(range(ci, m.start()))
                ci = m.end()
            kept_indices.extend(range(ci, len(joined)))
            # 按保留字符重新组合 runs（相邻字符 bold/footnote_ref 相同则合并）
            new_runs = []
            for idx in kept_indices:
                ch = joined[idx]
                bold, fn_ref = char_attrs[idx]
                if new_runs and new_runs[-1]['bold'] == bold and new_runs[-1]['footnote_ref'] == fn_ref:
                    new_runs[-1] = {**new_runs[-1], 'text': new_runs[-1]['text'] + ch}
                else:
                    new_runs.append({'text': ch, 'bold': bold, 'footnote_ref': fn_ref})
            runs_data = new_runs
        para_text = cleaned

        if not para_text.strip():
            continue

        if para_text in seen_texts:
            continue
        seen_texts.add(para_text)

        result.append({
            'runs': runs_data,
            'para_text': para_text,
            'has_triangle': has_triangle,
        })

    return result


def _is_run_element_bold(run_el, para_style_bold: bool = False) -> bool:
    """
    判断 lxml run 元素是否为粗体。检测顺序：
    1. rPr 中有 w:b：显式设置，直接用其值（val=0/false 则非粗体）
    2. rPr 中 w:rFonts[w:cs] 或 [w:eastAsia] 为黑体：视为粗体
    3. 均无显式设置：继承 para_style_bold（段落样式的粗体设置）
    """
    rpr = run_el.find(w('rPr'))
    if rpr is None:
        return para_style_bold

    # 显式 w:b
    b_el = rpr.find(w('b'))
    if b_el is not None:
        val = b_el.get(w('val'), '1')
        return val not in ('0', 'false')

    # 字体为黑体视为粗体
    rf = rpr.find(w('rFonts'))
    if rf is not None:
        cs_font = rf.get(w('cs'), '')
        ea_font = rf.get(w('eastAsia'), '')
        if '黑体' in cs_font or '黑体' in ea_font:
            return True

    return para_style_bold


def extract_footnotes(docx_path: str) -> dict:
    """
    从研报源 footnotes.xml 提取所有正文脚注文本。
    跳过 type=separator/continuationSeparator 的系统脚注（不论 id 数值）。
    返回 {src_id: url_text} 字典，src_id 为源文件中的整数编号。
    """
    with zipfile.ZipFile(docx_path) as z:
        if 'word/footnotes.xml' not in z.namelist():
            return {}
        fn_content = z.read('word/footnotes.xml')

    fn_tree = etree.fromstring(fn_content)
    result = {}
    SKIP_TYPES = {'separator', 'continuationSeparator'}

    for fn in fn_tree.findall('.//' + w('footnote')):
        fn_type = fn.get(w('type'))
        if fn_type in SKIP_TYPES:
            continue
        fn_id_str = fn.get(w('id'))
        if fn_id_str is None:
            continue
        fn_id = int(fn_id_str)
        texts = [t.text for t in fn.findall('.//' + w('t')) if t.text]
        # 去掉脚注引用符号本身（footnoteRef run 没有文本，只有第一个 run 的 t 会是空格+url）
        text = ''.join(texts).strip()
        if text:
            result[fn_id] = text

    return result


def extract_all(docx_path: str) -> dict:
    """一次性提取研报源所有信息，返回字典。"""
    highlighted = extract_highlighted_paragraphs(docx_path)
    abstract_text = highlighted[0]['para_text'] if highlighted else ''
    # footnotes: {src_id: url_text}，src_id 为源文件整数编号（可从 0 开始）
    footnotes_dict = extract_footnotes(docx_path)

    return {
        'date': extract_date(docx_path),
        'institution': extract_institution(docx_path),
        'title': extract_title(docx_path),
        'authors': extract_authors(docx_path),
        'highlighted_paragraphs': highlighted,
        'footnotes': footnotes_dict,
        'abstract_text': abstract_text,
    }
