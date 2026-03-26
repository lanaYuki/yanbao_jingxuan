"""
builder.py - 生成研报精选 docx
按照格式规则.md 中的规格构建文档。
包含：封面目录页、正文、脚注分割线、页眉logo。
"""

import os
import copy
import zipfile
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

FONT_HEITI     = '黑体'
FONT_FANG_SONG = '仿宋_GB2312'
FONT_HUA_ZHONG = '华文中宋'
FONT_TIMES     = 'Times New Roman'

SIZE_SECTION  = Pt(16)
SIZE_TITLE    = Pt(16)
SIZE_BODY     = Pt(14)
SIZE_BODY_EN  = Pt(14)   # 正文英文/数字：四号 = 14pt
SIZE_CITATION = Pt(13)
SIZE_COVER    = Pt(16)   # 封面目录页字号
SIZE_COVER_EN = Pt(16)   # 封面英文/数字：四号 = 16pt（封面用16pt对应四号）
SIZE_FOOTNOTE = Pt(10.5) # 脚注：Times New Roman 五号 = 10.5pt

LOGO_PATH = os.path.join(os.path.dirname(__file__), 'template_logo.png')
TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), 'CGI每周研报精选（第259期）-摘要.docx')

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
WP_NS = 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
A_NS  = 'http://schemas.openxmlformats.org/drawingml/2006/main'
PIC_NS = 'http://schemas.openxmlformats.org/drawingml/2006/picture'
R_NS  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
REL_IMAGE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image'


# ── 工具函数 ───────────────────────────────────────────


def _w(tag): return f'{{{W_NS}}}{tag}'


def _set_para_format(para, line_spacing_pt, space_after_pt=None, space_before_pt=None):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(line_spacing_pt)
    if space_after_pt is not None:
        pf.space_after = Pt(space_after_pt)
    if space_before_pt is not None:
        pf.space_before = Pt(space_before_pt)


def _set_run_font(run, font_name: str, size: Pt, bold: bool = False):
    """设置 run 的中英文字体、字号、粗体。"""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    rpr = run._r.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)


def _set_run_font_mixed(run, cn_font: str, en_font: str, cn_size: Pt, en_size: Pt, bold: bool = False):
    """
    设置 run 字体：中文用 cn_font/cn_size，英文/数字用 en_font/en_size。
    通过 w:rFonts 分别设置 ascii/hAnsi（英文）和 eastAsia（中文）。
    """
    run.bold = bold
    rpr = run._r.get_or_add_rPr()

    # 字体
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), cn_font)
    rFonts.set(qn('w:ascii'), en_font)
    rFonts.set(qn('w:hAnsi'), en_font)

    # 字号：Word 用中文字号（eastAsiaSize 通过 sz 和 szCs 控制）
    # sz 控制西文，szCs 控制复杂文字，eastAsia 跟 sz 走
    # 我们用两个 sz：中文字号放在 szCs，西文字号放在 sz
    # 实际上 Word 的中英混排：sz 同时影响西文和东亚字符
    # 正确做法：用 sz 设中文字号，英文字号通过分拆 run 实现（见 _split_mixed_runs）
    run.font.size = cn_size
    sz = rpr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rpr.append(sz)
    sz.set(qn('w:val'), str(int(cn_size.pt * 2)))
    szCs = rpr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rpr.append(szCs)
    szCs.set(qn('w:val'), str(int(cn_size.pt * 2)))


def _is_ascii_char(ch: str) -> bool:
    """判断单个字符是否为英文字母或数字（需要用 Times New Roman）。"""
    return ch.isascii() and (ch.isalpha() or ch.isdigit())


def _split_mixed_runs(text: str, bold: bool, cn_font: str, en_font: str,
                       cn_size: Pt, en_size: Pt) -> list:
    """
    将文本按中文/英文数字拆分，返回 run 参数列表：
    [{'text': str, 'font': str, 'size': Pt, 'bold': bool}, ...]
    """
    if not text:
        return []

    segments = []
    current = text[0]
    current_is_en = _is_ascii_char(text[0])

    for ch in text[1:]:
        ch_is_en = _is_ascii_char(ch)
        if ch_is_en == current_is_en:
            current += ch
        else:
            segments.append((current, current_is_en))
            current = ch
            current_is_en = ch_is_en
    segments.append((current, current_is_en))

    result = []
    for seg_text, is_en in segments:
        result.append({
            'text': seg_text,
            'font': en_font if is_en else cn_font,
            'size': en_size if is_en else cn_size,
            'bold': bold,
        })
    return result


def _add_run_to_para(para, text: str, font: str, size: Pt, bold: bool = False,
                     superscript: bool = False):
    """向段落添加一个已设置好字体/字号的 run，可选上标。"""
    run = para.add_run(text)
    run.bold = bold
    rpr = run._r.get_or_add_rPr()

    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)

    sz_val = str(int(size.pt * 2))
    sz = rpr.find(qn('w:sz'))
    if sz is None:
        sz = OxmlElement('w:sz')
        rpr.append(sz)
    sz.set(qn('w:val'), sz_val)
    szCs = rpr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rpr.append(szCs)
    szCs.set(qn('w:val'), sz_val)

    if superscript:
        vertAlign = OxmlElement('w:vertAlign')
        vertAlign.set(qn('w:val'), 'superscript')
        rpr.append(vertAlign)

    return run


def _add_hyperlink_run(para, text: str, url: str, font: str, size: Pt, bold: bool = False):
    """
    在段落中插入带超链接的 run。
    通过操作 XML 直接添加 w:hyperlink 元素。
    """
    # 先在文档关系中注册超链接
    part = para.part
    rId = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
                         is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), rId)

    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 字体
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rPr.append(rFonts)

    # 字号
    sz_val = str(int(size.pt * 2))
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), sz_val)
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), sz_val)
    rPr.append(szCs)

    # 下划线（超链接标准样式）
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # 蓝色（超链接标准颜色）
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    run_el.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    if text.startswith(' ') or text.endswith(' '):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run_el.append(t)

    hyperlink.append(run_el)
    para._p.append(hyperlink)


def _add_footnote_ref_to_para(para, fn_new_id: int):
    """
    在段落中插入脚注上标引用（w:footnoteReference）。
    使用 Word 内置的脚注引用字符样式（FootnoteReference / a0），
    让 Word 自动处理上标位置，避免手动 vertAlign 导致位置偏低。
    """
    run_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # 使用 Word 内置脚注引用样式，Word 会自动上标到右上角
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'a0')  # FootnoteReference 的内置样式 ID
    rPr.append(rStyle)
    run_el.append(rPr)

    fn_ref = OxmlElement('w:footnoteReference')
    fn_ref.set(qn('w:id'), str(fn_new_id))
    run_el.append(fn_ref)

    para._p.append(run_el)


def _set_para_indent_zero(para):
    """消除段落首行缩进，用于封面目录段落。"""
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')


def _add_page_break(doc: Document):
    """在文档末尾添加分页符（新起一页）。"""
    para = doc.add_paragraph()
    run = para.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)
    # 段落格式清零
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


# ── 封面目录页 ─────────────────────────────────────────


def _add_cover_page(doc: Document, issue: str, issue_date: str, articles: list):
    """
    生成封面目录页（第一页），格式参考精选模板 sdt 块。
    结构：
      日期+期号（右对齐，黑体，10pt）
      [对每篇文章：栏目名（黑体，16pt，加粗）/ 标题（黑体，16pt，加粗）/ 摘要（16pt，不加粗）]
    最后追加分页符。
    """
    # 日期行：右对齐，黑体，10pt，行距固定12pt
    date_para = doc.add_paragraph()
    date_para.alignment = 2  # RIGHT
    pf = date_para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(12)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    _set_para_indent_zero(date_para)
    date_run = date_para.add_run(f'{issue_date}　第{issue}期' if issue else issue_date)
    _set_run_font(date_run, FONT_HEITI, Pt(10), bold=False)

    # 每篇文章的目录条目
    seen_sections = set()
    for article in articles:
        section = article.get('section', '')
        title   = article.get('title', '')
        summary = article.get('summary', '')  # AI生成/用户编辑的摘要

        # 栏目名（新栏目才显示）
        if section and section not in seen_sections:
            seen_sections.add(section)
            p_sec = doc.add_paragraph()
            _set_para_indent_zero(p_sec)
            pf = p_sec.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(24)
            pf.space_before = Pt(8)
            pf.space_after = Pt(6)
            run = p_sec.add_run(section)
            _set_run_font(run, FONT_HEITI, SIZE_COVER, bold=True)
            run.underline = True

        # 篇目标题
        p_title = doc.add_paragraph()
        _set_para_indent_zero(p_title)
        pf = p_title.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(24)
        pf.space_before = Pt(0)
        pf.space_after = Pt(6)
        run_t = p_title.add_run(title)
        _set_run_font(run_t, FONT_HEITI, SIZE_COVER, bold=True)

        # 摘要文字（AI生成或用户填写），中英文混排
        if summary:
            p_sum = doc.add_paragraph()
            _set_para_indent_zero(p_sum)
            pf = p_sum.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(24)
            pf.space_before = Pt(0)
            pf.space_after = Pt(8)
            for seg in _split_mixed_runs(summary, bold=False,
                                          cn_font=FONT_FANG_SONG, en_font=FONT_TIMES,
                                          cn_size=SIZE_COVER, en_size=SIZE_COVER):
                _add_run_to_para(p_sum, seg['text'], seg['font'], seg['size'], seg['bold'])

    # 分页符（封面页结束）
    _add_page_break(doc)


# ── 正文段落 ───────────────────────────────────────────


def add_section_heading(doc: Document, text: str):
    """栏目名：黑体，16pt，加粗，下划线，段前14pt，段后12pt，行距固定24磅。"""
    para = doc.add_paragraph()
    _set_para_format(para, line_spacing_pt=24, space_after_pt=12, space_before_pt=14)
    run = para.add_run(text)
    _set_run_font(run, FONT_HEITI, SIZE_SECTION, bold=True)
    run.underline = True


def add_article_title(doc: Document, text: str):
    """篇目标题：黑体，16pt，加粗，段后12pt，行距固定24磅。"""
    para = doc.add_paragraph()
    _set_para_format(para, line_spacing_pt=24, space_after_pt=12)
    run = para.add_run(text)
    _set_run_font(run, FONT_HEITI, SIZE_TITLE, bold=True)


def add_body_paragraph(doc: Document, runs_data: list, line_spacing_pt: float = 24,
                        has_triangle: bool = False, footnote_id_map: dict = None,
                        link_map: dict = None):
    """
    正文段落：仿宋_GB2312，14pt，段后12pt，行距固定（默认24磅）。
    - has_triangle: 若为 True，段首插入 ► 加半角空格（不加全角缩进）
    - 普通段落：段首加两个全角空格缩进
    - 英文/数字用 Times New Roman 四号，中文用仿宋
    - footnote_id_map: {源脚注id: 精选新编号}，用于插入脚注上标
    - link_map: {文章名: url}，用于替换文中文章名为超链接
    """
    para = doc.add_paragraph()
    _set_para_format(para, line_spacing_pt=line_spacing_pt, space_after_pt=6)

    if has_triangle:
        # 两个全角空格缩进 + ► + 半角空格
        _add_run_to_para(para, '\u3000\u3000► ', FONT_FANG_SONG, SIZE_BODY, bold=False)
    else:
        # 首行空两格（全角空格）
        _add_run_to_para(para, '\u3000\u3000', FONT_FANG_SONG, SIZE_BODY, bold=False)

    if footnote_id_map is None:
        footnote_id_map = {}
    if link_map is None:
        link_map = {}

    for run_info in runs_data:
        text = run_info['text']
        bold = run_info['bold']
        fn_src_id = run_info.get('footnote_ref')  # 源文件脚注编号（可能为 None）

        # 超链接替换：检查文本中是否包含 link_map 中的文章名
        # 按文章名长度降序匹配，避免短名称覆盖长名称
        if link_map:
            text = _apply_link_map_to_para(para, text, bold, link_map)
        else:
            # 无超链接，直接混排输出
            for seg in _split_mixed_runs(text, bold=bold,
                                          cn_font=FONT_FANG_SONG, en_font=FONT_TIMES,
                                          cn_size=SIZE_BODY, en_size=SIZE_BODY_EN):
                _add_run_to_para(para, seg['text'], seg['font'], seg['size'], seg['bold'])

        # 插入脚注上标（紧跟在本 run 后），fn_src_id 可以为 0
        if fn_src_id is not None and fn_src_id in footnote_id_map:
            _add_footnote_ref_to_para(para, footnote_id_map[fn_src_id])


def _apply_link_map_to_para(para, text: str, bold: bool, link_map: dict) -> str:
    """
    在 para 中输出 text，同时将匹配到 link_map 中文章名的部分替换为超链接 run。
    返回空字符串（文本已全部输出到 para）。
    """
    # 按文章名长度降序，避免短名遮盖长名
    names_sorted = sorted(link_map.keys(), key=len, reverse=True)
    # 用正则将文本切分为：普通文本段 / 匹配段
    import re
    pattern = '(' + '|'.join(re.escape(n) for n in names_sorted) + ')'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part in link_map:
            _add_hyperlink_run(para, part, link_map[part],
                               font=FONT_FANG_SONG, size=SIZE_BODY, bold=bold)
        else:
            for seg in _split_mixed_runs(part, bold=bold,
                                          cn_font=FONT_FANG_SONG, en_font=FONT_TIMES,
                                          cn_size=SIZE_BODY, en_size=SIZE_BODY_EN):
                _add_run_to_para(para, seg['text'], seg['font'], seg['size'], seg['bold'])
    return ''


def _count_para_lines(text: str, chars_per_line: int = 31,
                       line_spacing_pt: float = 24, space_after_pt: float = 6) -> float:
    """
    估算段落占用的等效行数（单位：line_spacing_pt 的行高）。
    含首行两个全角空格缩进，并将 space_after 折算为行高单位。
    """
    import math
    length = len(text) + 2  # +2 for 全角空格缩进
    text_lines = math.ceil(max(length, 1) / chars_per_line)
    # space_after 折算为行高单位
    extra = space_after_pt / line_spacing_pt
    return text_lines + extra


def add_citation_section(doc: Document, date: str, institution: str,
                          title: str, authors: list,
                          body_lines_used: float,
                          footnote_count: int = 0):
    """
    归属声明区：长横线 + 来源行 + 作者行，压底显示。
    body_lines_used：正文已用行数（精确计算，单位为24pt行）。
    footnote_count：本页脚注数量，用于从可用页高中扣除脚注区占用的高度。
    规则：
    - 总行数 = 1（来源行）+ len(authors)
    - ≤4行：当前页压底，全部在同页
    - >4行：当前页压底只放来源行，分页后放所有作者行（保留横线，不压底）
    """
    # 可用页高：A4(29.7cm) - 上边距(3.2cm) - 下边距(2.2cm) = 24.3cm
    # 每行高约 24pt；因此每页约 28.7 行（取28行保守）
    PAGE_LINES = 28.0

    # 脚注区占用高度：separator≈6pt + 每条脚注≈14pt（10.5pt字+间距）
    # 折算为 24pt 等效行数，从可用页高中扣除
    if footnote_count > 0:
        fn_area_pt = 6 + footnote_count * 14   # separator + 脚注行
        PAGE_LINES -= fn_area_pt / 24.0

    total_info_lines = 1 + len(authors)  # 来源行 + 作者行

    def _add_separator(doc):
        """长横线：段落上边框"""
        para = doc.add_paragraph()
        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(6)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '6')
        top.set(qn('w:space'), '1')
        top.set(qn('w:color'), '000000')
        pBdr.append(top)
        pPr.append(pBdr)

    def _add_source_line(doc):
        p = doc.add_paragraph()
        _set_para_format(p, line_spacing_pt=18, space_before_pt=4, space_after_pt=0)
        run = p.add_run(f'以上观点来自：{date}{institution}已发布的《{title}》')
        _set_run_font(run, FONT_FANG_SONG, SIZE_CITATION, bold=True)

    def _add_author_lines(doc):
        for author in authors:
            p = doc.add_paragraph()
            _set_para_format(p, line_spacing_pt=18, space_after_pt=0)
            sfc_part = f' SFC CE Ref：{author["sfc"]}' if author.get('sfc') else ''
            text = f'{author["name"]} 分析员 SAC 执业证书编号：{author["sac"]}{sfc_part}'
            run = p.add_run(text)
            _set_run_font(run, FONT_FANG_SONG, SIZE_CITATION, bold=False)

    def _add_fill(doc, used_lines, info_lines):
        """插入空白段落使内容压底（每个空白段落 = 1行24pt）"""
        fill = PAGE_LINES - used_lines - info_lines - 1  # -1 for separator
        count = max(0, round(fill))
        for _ in range(count):
            spacer = doc.add_paragraph()
            pf = spacer.paragraph_format
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(24)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)

    if total_info_lines <= 4:
        _add_fill(doc, body_lines_used, total_info_lines)
        _add_separator(doc)
        _add_source_line(doc)
        _add_author_lines(doc)
    else:
        # 来源行压底在当前页
        _add_fill(doc, body_lines_used, 2)  # 横线 + 来源行
        _add_separator(doc)
        _add_source_line(doc)
        # 换页，新页放作者行（保留横线，不压底）
        _add_page_break(doc)
        _add_separator(doc)
        _add_author_lines(doc)


# ── 脚注 ──────────────────────────────────────────────


def _add_footnotes_to_docx(output_path: str, all_footnotes: list[str]):
    """
    将所有文章的脚注写入 docx 的 footnotes.xml。
    包含标准的 separator（长横线）和 continuationSeparator。
    脚注从 id=1 开始重新编号。
    此函数直接操作 zip 文件。
    """
    if not all_footnotes:
        return

    SEP_ID = -1
    CONT_ID = 0

    fn_entries = []

    # separator（脚注区分割线，Word 标准格式）
    fn_entries.append(f'''<w:footnote w:type="separator" w:id="{SEP_ID}">
  <w:p>
    <w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
    <w:r><w:separator/></w:r>
  </w:p>
</w:footnote>''')

    # continuationSeparator
    fn_entries.append(f'''<w:footnote w:type="continuationSeparator" w:id="{CONT_ID}">
  <w:p>
    <w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
    <w:r><w:continuationSeparator/></w:r>
  </w:p>
</w:footnote>''')

    # 正文脚注：上标数字 + 脚注文本，字体 Times New Roman 五号（10.5pt = 21 half-points）
    FN_SZ = '21'  # 10.5pt * 2
    FN_FONT = 'Times New Roman'
    for idx, url in enumerate(all_footnotes, start=1):
        fn_entries.append(f'''<w:footnote w:id="{idx}">
  <w:p>
    <w:pPr><w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/></w:pPr>
    <w:r>
      <w:rPr>
        <w:rFonts w:ascii="{FN_FONT}" w:hAnsi="{FN_FONT}" w:eastAsia="{FN_FONT}"/>
        <w:sz w:val="{FN_SZ}"/><w:szCs w:val="{FN_SZ}"/>
        <w:vertAlign w:val="superscript"/>
      </w:rPr>
      <w:t>{idx}</w:t>
    </w:r>
    <w:r>
      <w:rPr>
        <w:rFonts w:ascii="{FN_FONT}" w:hAnsi="{FN_FONT}" w:eastAsia="{FN_FONT}"/>
        <w:sz w:val="{FN_SZ}"/><w:szCs w:val="{FN_SZ}"/>
      </w:rPr>
      <w:t xml:space="preserve"> {url}</w:t>
    </w:r>
  </w:p>
</w:footnote>''')

    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        '<w:footnotes xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">\n'
        + '\n'.join(fn_entries)
        + '\n</w:footnotes>'
    )

    REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
    FN_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
    CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
    FN_CT = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'

    # 将 footnotes.xml 写入 docx（重新打包 zip）
    buf = BytesIO()
    with zipfile.ZipFile(output_path, 'r') as zin:
        existing = {i.filename for i in zin.infolist()}
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)

                if item.filename == 'word/footnotes.xml':
                    # 覆盖已有的 footnotes.xml
                    zout.writestr(item, footnotes_xml)
                    continue

                if item.filename == 'word/_rels/document.xml.rels':
                    # 注册 footnotes.xml 关系（若尚未存在）
                    tree = etree.fromstring(data)
                    has_fn_rel = any(
                        r.get('Type') == FN_REL_TYPE
                        for r in tree.findall(f'{{{REL_NS}}}Relationship')
                    )
                    if not has_fn_rel:
                        rel = etree.SubElement(tree, f'{{{REL_NS}}}Relationship')
                        rel.set('Id', 'rId_fn1')
                        rel.set('Type', FN_REL_TYPE)
                        rel.set('Target', 'footnotes.xml')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                elif item.filename == '[Content_Types].xml':
                    # 注册 footnotes.xml 的 ContentType（若尚未存在）
                    tree = etree.fromstring(data)
                    has_fn_ct = any(
                        o.get('PartName') == '/word/footnotes.xml'
                        for o in tree.findall(f'{{{CT_NS}}}Override')
                    )
                    if not has_fn_ct:
                        ov = etree.SubElement(tree, f'{{{CT_NS}}}Override')
                        ov.set('PartName', '/word/footnotes.xml')
                        ov.set('ContentType', FN_CT)
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                zout.writestr(item, data)

            # 如果原 docx 没有 footnotes.xml，新增文件
            if 'word/footnotes.xml' not in existing:
                zout.writestr('word/footnotes.xml', footnotes_xml)

    with open(output_path, 'wb') as f:
        f.write(buf.getvalue())


# ── 页眉 Logo ─────────────────────────────────────────


def _add_header_to_docx(output_path: str):
    """
    从模板 docx 复制页眉到输出 docx：
    - 第一页（first）：header1.xml（大Logo）
    - 其他页（default）：header4.xml（"研报精选"文字 + 灰色线下移 + 紧贴对齐）
    """
    if not os.path.exists(TEMPLATE_PATH):
        return
    
    # ── 参数配置 ──
    # 增加此值使灰色线下移（360000 EMU = 1cm）。
    # 如果觉得位置还不够靠下，可将 250000 继续调大。
    MOVE_DOWN_EMU = 120000 

    REL_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
    W_NS_LOCAL = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R_NS_LOCAL = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    WPS_NS = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    A_NS_H = 'http://schemas.openxmlformats.org/drawingml/2006/main'

    with zipfile.ZipFile(TEMPLATE_PATH, 'r') as tmpl:
        header1_xml = tmpl.read('word/header1.xml')
        header1_rels = tmpl.read('word/_rels/header1.xml.rels')
        header4_xml = tmpl.read('word/header4.xml')
        header4_rels = tmpl.read('word/_rels/header4.xml.rels')
        image1_data = tmpl.read('word/media/image1.png')
        image2_data = tmpl.read('word/media/image2.png')

    h4_tree = etree.fromstring(header4_xml)

    # 1. 调整文本框 (wsp) 的容器位置和文字对齐
    for wsp in h4_tree.iter(f'{{{WPS_NS}}}wsp'):
        for xfrm in wsp.iter(f'{{{A_NS_H}}}xfrm'):
            off = xfrm.find(f'{{{A_NS_H}}}off')
            ext = xfrm.find(f'{{{A_NS_H}}}ext')
            if off is not None and ext is not None:
                # 设置文本框宽度为全宽 (约18.6cm)，确保不会挤压文字
                ext.set('cx', '6726754') 
                # x=0 靠左对齐；y 坐标决定文字垂直位置
                # 调小 1850000 这个基数会使文字更向上“贴近”灰色线
                off.set('x', '0')
                off.set('y', str(2050000 + MOVE_DOWN_EMU))
            break 

    # 2. 调整灰色线 (Shape) 的垂直位置
    # 查找所有的 xfrm，排除属于文本框内部的，只针对线形状
    for xfrm in h4_tree.xpath("//a:xfrm", namespaces={'a': A_NS_H}):
        if f'{{{WPS_NS}}}wsp' not in [anc.tag for anc in xfrm.iterancestors()]:
            off = xfrm.find(f'{{{A_NS_H}}}off')
            if off is not None:
                current_y = int(off.get('y', '0'))
                off.set('y', str(current_y + MOVE_DOWN_EMU))

    # 3. 彻底清除文本框内部段落的缩进和边距
    for txbx_para in h4_tree.findall(f'.//{{{W_NS_LOCAL}}}txbxContent/{{{W_NS_LOCAL}}}p'):
        pPr = txbx_para.find(f'{{{W_NS_LOCAL}}}pPr')
        if pPr is None:
            pPr = etree.SubElement(txbx_para, f'{{{W_NS_LOCAL}}}pPr')
            txbx_para.insert(0, pPr)
        
        # 强制右对齐
        jc = pPr.find(f'{{{W_NS_LOCAL}}}jc') or etree.SubElement(pPr, f'{{{W_NS_LOCAL}}}jc')
        jc.set(f'{{{W_NS_LOCAL}}}val', 'right')
        
        # 移除缩进 (ind)
        ind = pPr.find(f'{{{W_NS_LOCAL}}}ind')
        if ind is not None: pPr.remove(ind)
        new_ind = etree.SubElement(pPr, f'{{{W_NS_LOCAL}}}ind')
        new_ind.set(f'{{{W_NS_LOCAL}}}left', '0')
        new_ind.set(f'{{{W_NS_LOCAL}}}firstLine', '0')

        # 移除段前段后间距，使文字紧贴灰色线
        spacing = pPr.find(f'{{{W_NS_LOCAL}}}spacing') or etree.SubElement(pPr, f'{{{W_NS_LOCAL}}}spacing')
        spacing.set(f'{{{W_NS_LOCAL}}}before', '0')
        spacing.set(f'{{{W_NS_LOCAL}}}after', '0')

        # 设置字体大小为 14pt (28)
        for run in txbx_para.findall(f'{{{W_NS_LOCAL}}}r'):
            rPr = run.find(f'{{{W_NS_LOCAL}}}rPr') or etree.SubElement(run, f'{{{W_NS_LOCAL}}}rPr')
            sz = rPr.find(f'{{{W_NS_LOCAL}}}sz') or etree.SubElement(rPr, f'{{{W_NS_LOCAL}}}sz')
            sz.set(f'{{{W_NS_LOCAL}}}val', '28')

    header4_xml = etree.tostring(h4_tree, xml_declaration=True, encoding='UTF-8', standalone=True)

    # ── 写入输出文件 ──
    buf = BytesIO()
    with zipfile.ZipFile(output_path, 'r') as zin:
        with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                
                # 更新文档关系
                if item.filename == 'word/_rels/document.xml.rels':
                    tree = etree.fromstring(data)
                    for rid, target in [('rId100', 'header1.xml'), ('rId101', 'header4.xml')]:
                        rel = etree.SubElement(tree, f'{{{REL_NS}}}Relationship')
                        rel.set('Id', rid)
                        rel.set('Type', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/header')
                        rel.set('Target', target)
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                # 在主文档中应用页眉
                elif item.filename == 'word/document.xml':
                    tree = etree.fromstring(data)
                    sectPr = tree.find(f'.//{{{W_NS_LOCAL}}}sectPr')
                    if sectPr is not None:
                        for old in sectPr.findall(f'{{{W_NS_LOCAL}}}headerReference'):
                            sectPr.remove(old)
                        if sectPr.find(f'{{{W_NS_LOCAL}}}titlePg') is None:
                            sectPr.insert(0, etree.Element(f'{{{W_NS_LOCAL}}}titlePg'))
                        h1 = etree.Element(f'{{{W_NS_LOCAL}}}headerReference')
                        h1.set(f'{{{W_NS_LOCAL}}}type', 'first')
                        h1.set(f'{{{R_NS_LOCAL}}}id', 'rId100')
                        sectPr.insert(0, h1)
                        h4 = etree.Element(f'{{{W_NS_LOCAL}}}headerReference')
                        h4.set(f'{{{W_NS_LOCAL}}}type', 'default')
                        h4.set(f'{{{R_NS_LOCAL}}}id', 'rId101')
                        sectPr.insert(0, h4)
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                # 更新内容类型定义
                elif item.filename == '[Content_Types].xml':
                    tree = etree.fromstring(data)
                    CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
                    for p in ['/word/header1.xml', '/word/header4.xml']:
                        ov = etree.SubElement(tree, f'{{{CT_NS}}}Override')
                        ov.set('PartName', p)
                        ov.set('ContentType', 'application/vnd.openxmlformats-officedocument.wordprocessingml.header+xml')
                    data = etree.tostring(tree, xml_declaration=True, encoding='UTF-8', standalone=True)

                zout.writestr(item, data)

            # 写入页眉 XML 和 媒体文件
            zout.writestr('word/header1.xml', header1_xml)
            zout.writestr('word/_rels/header1.xml.rels', header1_rels)
            zout.writestr('word/header4.xml', header4_xml)
            zout.writestr('word/_rels/header4.xml.rels', header4_rels)
            zout.writestr('word/media/image1.png', image1_data)
            zout.writestr('word/media/image2.png', image2_data)

    with open(output_path, 'wb') as f:
        f.write(buf.getvalue())


# ── 主构建函数 ────────────────────────────────────────


def build_jingxuan(articles: list, output_path: str, issue: str = '',
                    link_map: dict = None, list_categories: dict = None):
    """
    构建研报精选 docx。

    articles: 列表，每个元素字典：
      {
        'section': str,
        'title': str,
        'paragraphs': list,   # extract_highlighted_paragraphs 结果
        'date': str,
        'institution': str,
        'authors': list,
        'summary': str,       # AI或用户摘要（封面目录用）
        'footnotes': list,    # 本篇脚注 URL 列表（按源文件编号顺序）
        'footnote_refs': dict,# {源脚注id: url} 映射（由 extractor 提供）
      }
    link_map: {文章名: url}，用于替换正文中的文章名为超链接。
    """
    if link_map is None:
        link_map = {}

    doc = Document()

    # 页面设置：上3.2cm，下/左/右2.2cm
    for section in doc.sections:
        section.top_margin = Cm(3.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # 删除默认空段落
    for para in doc.paragraphs:
        para._element.getparent().remove(para._element)

    # 推断封面日期：取第一篇的 date，若无则空
    cover_date = articles[0].get('date', '') if articles else ''

    # ── 封面目录页 ──
    _add_cover_page(doc, issue, cover_date, articles)

    # ── 构建全局脚注列表（所有文章的脚注合并，重新编号）──
    # 结构：all_footnote_urls[new_id-1] = url（new_id 从 1 开始）
    # footnote_id_maps[article_idx] = {src_id: new_id}
    all_footnote_urls = []
    article_fn_maps = []  # per-article {src_id: new_global_id}

    for article in articles:
        paragraphs = article.get('paragraphs', [])
        # footnotes 现在是 {src_id: url} 字典，src_id 可以从 0 开始
        src_footnotes = article.get('footnotes', {})
        fn_map = {}  # src_id -> new_global_id

        # 收集本篇段落中实际引用到的脚注 src_id（包含 id=0）
        used_src_ids = set()
        for para_data in paragraphs:
            for run_info in para_data.get('runs', []):
                src_id = run_info.get('footnote_ref')
                if src_id is not None:  # 允许 id=0
                    used_src_ids.add(src_id)

        # 按 src_id 从小到大分配新编号（全局连续编号从1开始）
        for src_id in sorted(used_src_ids):
            if src_id in src_footnotes:
                url = src_footnotes[src_id]
                new_id = len(all_footnote_urls) + 1
                all_footnote_urls.append(url)
                fn_map[src_id] = new_id

        article_fn_maps.append(fn_map)

    # ── 正文（每篇研报）──
    from ai_helper import check_and_fix_orphan_lines

    for i, article in enumerate(articles):
        section_name = article.get('section', '')
        title        = article.get('title', '')
        paragraphs   = article.get('paragraphs', [])
        date         = article.get('date', '')
        institution  = article.get('institution', '中金公司研究部')
        authors      = article.get('authors', [])
        fn_map       = article_fn_maps[i]

        # ── AI孤行检查 ──
        # 页高：A4(29.7cm) - 上(3.2cm) - 下(2.2cm) = 24.3cm；24pt行≈0.847cm → ~28.7行
        PAGE_LINES = 28.0
        info_lines = 1 + len(authors)
        # 归属信息行高18pt + space_after约0 → 折算为24pt行：18/24 = 0.75
        info_lines_24pt = info_lines * 0.75
        # 标题行：1行24pt + space_after=12pt → 1 + 12/24 = 1.5 等效行
        header_lines = (1.5 if section_name else 0) + (1.5 if title else 0)
        avail_lines = PAGE_LINES - header_lines - info_lines_24pt - 1  # -1 for separator

        fix = check_and_fix_orphan_lines(
            paragraphs,
            page_lines=avail_lines,
            chars_per_line=33,
        )
        paragraphs = fix['paragraphs']
        body_line_spacing = fix['line_spacing']

        if section_name:
            add_section_heading(doc, section_name)
        if title:
            add_article_title(doc, title)
        for para_data in paragraphs:
            add_body_paragraph(
                doc,
                para_data['runs'],
                line_spacing_pt=body_line_spacing,
                has_triangle=para_data.get('has_triangle', False),
                footnote_id_map=fn_map,
                link_map=link_map,
            )
        if date or title:
            body_lines = float(header_lines)
            for para_data in paragraphs:
                body_lines += _count_para_lines(
                    para_data.get('para_text', ''),
                    line_spacing_pt=body_line_spacing,
                    space_after_pt=6,
                )
            add_citation_section(doc, date, institution, title, authors,
                                  body_lines_used=body_lines,
                                  footnote_count=len(fn_map))

        if i < len(articles) - 1:
            _add_page_break(doc)

    # ── 研报清单（附在研报精选之后）──
    if list_categories and any(list_categories.get(cat) for cat in
                                ['宏观', '策略及大宗商品', '固定收益', '行业']):
        from list_builder import append_list_to_doc
        append_list_to_doc(doc, list_categories)

    doc.save(output_path)

    # ── 写入脚注 ──
    if all_footnote_urls:
        _add_footnotes_to_docx(output_path, all_footnote_urls)

    # ── 写入页眉 Logo ──
    _add_header_to_docx(output_path)